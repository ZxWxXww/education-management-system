from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE = "基于 SpringBoot中小学教培机构管理系统的设计与实现"
TITLE_EN = "Design and Implementation of a Primary and Secondary Education Training Institution Management System Based on SpringBoot"
COLLEGE = "信息工程学院"
MAJOR = "软件工程"
STUDENT_NAME = "朱文欣"
STUDENT_ID = "22206091042"
ADVISOR = "童芳"
ADVISOR_TITLE = "高级工程师（硕士）"
COMPLETE_DATE = "2026 年 5 月 10 日"

ROOT = Path("D:/education")
MD_PATH = ROOT / "基于SpringBoot的中小学教培机构管理系统的设计与实现.md"
DOCX_PATH = ROOT / "基于 SpringBoot中小学教培机构管理系统的设计与实现.docx"


@dataclass
class Block:
    kind: str
    text: str


def set_document_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.left_margin = Cm(2.8)


def set_run_font(run, east="宋体", west="Times New Roman", size=12, bold=False) -> None:
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.font.bold = bold


def set_fixed_line_spacing(paragraph, points=20) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(points)
    fmt.first_line_indent = Pt(28)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_section_page_number(section, start=None, fmt=None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    if fmt is not None:
        pg_num.set(qn("w:fmt"), fmt)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("本科毕业论文")
    set_run_font(r, size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("( 2026 届 )")
    set_run_font(r, size=16)

    for label, value in [
        ("题    目：", TITLE),
        ("学    院：", COLLEGE),
        ("专    业：", MAJOR),
        ("学生姓名：", f"{STUDENT_NAME}    学号：{STUDENT_ID}"),
        ("指导教师：", f"{ADVISOR}    职称（学位）：{ADVISOR_TITLE}"),
        ("合作导师：", "                职称（学位）："),
        ("完成时间：", COMPLETE_DATE),
        ("成    绩：", ""),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(2.2)
        p.paragraph_format.space_before = Pt(10)
        r1 = p.add_run(label)
        set_run_font(r1, size=15)
        r2 = p.add_run(value)
        set_run_font(r2, size=15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("黄山学院教务处制")
    set_run_font(r, size=16)


def add_originality_statement(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("学位论文原创性声明")
    set_run_font(r, size=16, bold=True)

    body = (
        "兹呈交的学位论文，是本人在指导老师指导下独立完成的研究成果。"
        "本人在论文写作中参考的其他个人或集体的研究成果，均在文中以明确方式标明。"
        "本人依法享有和承担由此论文而产生的权利和责任。"
    )
    p = doc.add_paragraph()
    set_fixed_line_spacing(p)
    r = p.add_run(body)
    set_run_font(r, size=12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run("声明人（签名）：")
    set_run_font(r, size=12)

    p = doc.add_paragraph()
    r = p.add_run("年   月   日")
    set_run_font(r, size=12)


def add_toc(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目   录")
    set_run_font(r, size=22, bold=True)

    p = doc.add_paragraph()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "打开 Word 后右键目录并选择“更新域”，即可生成完整目录。"
    separate.append(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p._p.append(begin)
    p._p.append(instr)
    p._p.append(separate)
    p._p.append(end)


def parse_markdown(md_text: str) -> list[Block]:
    lines = md_text.splitlines()
    blocks: list[Block] = []
    buf: list[str] = []

    def flush() -> None:
        nonlocal buf
        if not buf:
            return
        text = " ".join(item.strip() for item in buf).strip()
        if text:
            blocks.append(Block("paragraph", text))
        buf = []

    started = False
    for raw in lines:
        line = raw.strip()
        if not started:
            if line == "## 摘要":
                started = True
                blocks.append(Block("h2", "摘要"))
            continue
        if not line:
            flush()
            continue
        if line.startswith("#### "):
            flush()
            blocks.append(Block("h4", line[5:].strip()))
            continue
        if line.startswith("### "):
            flush()
            blocks.append(Block("h3", line[4:].strip()))
            continue
        if line.startswith("## "):
            flush()
            blocks.append(Block("h2", line[3:].strip()))
            continue
        buf.append(line)
    flush()
    return blocks


def clean_inline_markdown(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text).strip()


def add_title_info_page(doc: Document, chinese: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE if chinese else TITLE_EN)
    set_run_font(r, east="宋体" if chinese else "Times New Roman", west="Times New Roman", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = f"{COLLEGE} {MAJOR} {STUDENT_NAME}（{STUDENT_ID}）" if chinese else f"{STUDENT_NAME} ({STUDENT_ID})"
    r = p.add_run(info)
    set_run_font(r, east="宋体" if chinese else "Times New Roman", west="Times New Roman", size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = f"指导老师：{ADVISOR}（{ADVISOR_TITLE}）" if chinese else f"Director: {ADVISOR}"
    r = p.add_run(info)
    set_run_font(r, east="宋体" if chinese else "Times New Roman", west="Times New Roman", size=14)

    if not chinese:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"({MAJOR}, {COLLEGE}, Huangshan University)")
        set_run_font(r, east="Times New Roman", west="Times New Roman", size=14)


def add_abstract_body(doc: Document, text: str, chinese: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("摘要：" if chinese else "Abstract: ")
    set_run_font(r1, east="黑体" if chinese else "Times New Roman", west="Times New Roman", size=10.5, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, east="宋体" if chinese else "Times New Roman", west="Times New Roman", size=10.5)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.first_line_indent = Pt(28)


def add_keyword(doc: Document, text: str, chinese: bool) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run("关键词：" if chinese else "Key Words: ")
    set_run_font(r1, east="黑体" if chinese else "Times New Roman", west="Times New Roman", size=10.5, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, east="宋体" if chinese else "Times New Roman", west="Times New Roman", size=10.5)
    p.paragraph_format.line_spacing = Pt(20)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=14 if level == 2 else 12, bold=True)
    if level == 2:
        p.style = doc.styles["Heading 1"]
    else:
        p.style = doc.styles["Heading 2"]


def add_body_paragraph(doc: Document, text: str) -> None:
    text = clean_inline_markdown(text)
    if text.startswith("表题："):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.replace("表题：", "", 1).strip())
        set_run_font(r, size=10.5)
        p.paragraph_format.line_spacing = Pt(20)
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=12)
    set_fixed_line_spacing(p)


def add_references(doc: Document, lines: list[str]) -> None:
    for line in lines:
        if not line:
            continue
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, size=10.5)
        fmt = p.paragraph_format
        fmt.line_spacing = Pt(20)
        fmt.left_indent = Pt(21)
        fmt.first_line_indent = Pt(-21)


def build_doc() -> None:
    blocks = parse_markdown(MD_PATH.read_text(encoding="utf-8"))

    doc = Document()
    enable_update_fields(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    add_cover(doc)
    add_originality_statement(doc)
    add_toc(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    abstract_section = doc.sections[-1]
    abstract_section.footer.is_linked_to_previous = False
    footer = abstract_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    set_section_page_number(abstract_section, start=1, fmt="upperRoman")

    idx = 0
    assert blocks[idx].text == "摘要"
    idx += 1

    zh_abs: list[str] = []
    zh_kw = ""
    while idx < len(blocks) and not (blocks[idx].kind == "h2" and blocks[idx].text == "Abstract"):
        current = clean_inline_markdown(blocks[idx].text)
        if "关键词：" in current:
            zh_kw = current.replace("关键词：", "").strip()
        elif blocks[idx].kind == "paragraph":
            zh_abs.append(current)
        idx += 1

    add_title_info_page(doc, chinese=True)
    add_abstract_body(doc, " ".join(zh_abs), chinese=True)
    add_keyword(doc, zh_kw, chinese=True)

    idx += 1
    en_abs: list[str] = []
    en_kw = ""
    while idx < len(blocks) and not (blocks[idx].kind == "h2" and re.match(r"1\s+", blocks[idx].text)):
        current = clean_inline_markdown(blocks[idx].text)
        if "Key Words:" in current:
            en_kw = current.replace("Key Words:", "").strip()
        elif blocks[idx].kind == "paragraph":
            en_abs.append(current)
        idx += 1

    doc.add_page_break()
    add_title_info_page(doc, chinese=False)
    add_abstract_body(doc, " ".join(en_abs), chinese=False)
    add_keyword(doc, en_kw, chinese=False)

    doc.add_section(WD_SECTION.NEW_PAGE)
    body_section = doc.sections[-1]
    body_section.footer.is_linked_to_previous = False
    footer = body_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    set_section_page_number(body_section, start=1, fmt="decimal")

    in_references = False
    ref_lines: list[str] = []
    for block in blocks[idx:]:
        if block.kind == "h2":
            if in_references and block.text != "致谢":
                add_references(doc, ref_lines)
                ref_lines.clear()
                in_references = False
            doc.add_page_break()
            add_heading(doc, block.text, 2)
            if block.text == "参考文献":
                in_references = True
            continue

        if in_references and block.kind == "paragraph":
            ref_lines.append(clean_inline_markdown(block.text))
            continue

        if block.kind in {"h3", "h4"}:
            add_heading(doc, block.text, 3)
        elif block.kind == "paragraph":
            add_body_paragraph(doc, block.text)

    if ref_lines:
        add_references(doc, ref_lines)

    set_document_margins(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
